"""Append clinical-context pilot screenshots to manuscript (no duplicate prior figures)."""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

PAPER = Path.home() / "OneDrive" / "Desktop" / "rbGyanX_mobile_paper"
MANUSCRIPT = PAPER / "rbGyanx_mobile_manuscript.docx"
FIGURES = PAPER / "figures" / "screenshots"
BACKUP = PAPER / f"rbGyanx_mobile_manuscript_backup_{datetime.now():%Y%m%d_%H%M%S}.docx"

FIGURES_TO_APPEND = [
    (
        "fig10_clinical_context_filled.png",
        "Figure. Clinical context form with investigator-entered H&N fields (age, HPV, smoking, chemotherapy).",
    ),
    (
        "fig11_results_clinical_context.png",
        "Figure. Calculation results Clinical tab showing entered context linked to TCP/NTCP output.",
    ),
]


def main() -> None:
    if not MANUSCRIPT.exists():
        raise SystemExit(f"Manuscript not found: {MANUSCRIPT}")

    shutil.copy2(MANUSCRIPT, BACKUP)
    doc = Document(str(MANUSCRIPT))

    doc.add_page_break()
    h = doc.add_heading("Clinical context (filled on device)", level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    inserted = 0
    for fname, caption in FIGURES_TO_APPEND:
        img = FIGURES / fname
        if not img.exists():
            print(f"SKIP missing {fname}")
            continue
        doc.add_picture(str(img), width=Inches(3.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        inserted += 1
        print(f"APPEND {fname}")

    doc.save(str(MANUSCRIPT))
    print(f"Saved {MANUSCRIPT} (+{inserted} figures)")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
