"""Embed pilot screenshots into rbGyanx_mobile_manuscript.docx."""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

PAPER = Path.home() / "OneDrive" / "Desktop" / "rbGyanX_mobile_paper"
MANUSCRIPT = PAPER / "rbGyanx_mobile_manuscript.docx"
FIGURES = PAPER / "figures" / "screenshots"
BACKUP = PAPER / f"rbGyanx_mobile_manuscript_backup_{datetime.now():%Y%m%d_%H%M%S}.docx"

FIGURE_ORDER = [
    ("fig01_app_home.png", "Figure. rbGyanX Mobile home screen (offline build v1.0.0)."),
    ("fig02_dvh_import.png", "Figure. Composite DVH import from device Downloads."),
    ("fig03_calculation_setup.png", "Figure. Plan calculation setup (HN, IMRT, prescription)."),
    ("fig04_calculation_results.png", "Figure. Composite TCP/NTCP results and dose metrics (RBX-TXT-001)."),
    ("fig05_results_RBX-TXT-004.png", "Figure. Second pilot case results (RBX-TXT-004, 50 Gy)."),
    ("fig06_therapeutic_window.png", "Figure. Therapeutic window visualization (TCP + NTCP vs dose)."),
    ("fig07_report_export.png", "Figure. PDF/DOCX report export screen."),
]


def main() -> None:
    if not MANUSCRIPT.exists():
        raise SystemExit(f"Manuscript not found: {MANUSCRIPT}")

    shutil.copy2(MANUSCRIPT, BACKUP)
    doc = Document(str(MANUSCRIPT))

    doc.add_page_break()
    h = doc.add_heading("Mobile app screenshots (pilot run)", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    inserted = 0
    for fname, caption in FIGURE_ORDER:
        img = FIGURES / fname
        if not img.exists():
            print(f"SKIP missing {fname}")
            continue
        doc.add_picture(str(img), width=Inches(3.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        inserted += 1
        print(f"INSERT {fname}")

    doc.save(str(MANUSCRIPT))
    print(f"Saved {MANUSCRIPT} ({inserted} figures)")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
